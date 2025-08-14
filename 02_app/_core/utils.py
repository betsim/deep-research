import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
from typing import List, Callable, Any, Optional
from tqdm import tqdm
import tiktoken
from datetime import datetime
import io
import re
from docx import Document
from docx.shared import RGBColor
from docx.oxml import parse_xml
from _core.config import config


class TokenCounter:
    """Handles token counting operations."""

    _encoders = {}

    def __init__(self):
        raise TypeError(
            "TokenCounter is a utility class and should not be instantiated."
        )

    @classmethod
    def count_tokens(cls, text: str, model: Optional[str] = None) -> int:
        """Count tokens in text using OpenAI's tokenizer."""
        model = model or config["llm"]["token_count_model"]

        # Cache encoders to avoid re-initialization
        if model not in cls._encoders:
            cls._encoders[model] = tiktoken.encoding_for_model(model)

        encoder = cls._encoders[model]

        try:
            tokens = encoder.encode(text)
            return len(tokens)
        except Exception as e:
            raise RuntimeError(f"Token counting failed: {e}") from e


def get_model_and_workflow_config(fast_mode=False):
    """Get model and workflow configuration based on mode"""
    if config["development"]["enabled"]:
        model_id = config["development"]["model_for_dev_testing"]
        model_config = {
            key: model_id
            for key in [
                "create_queries",
                "check_relevance",
                "analyze_documents",
                "reflect_task",
                "final_report",
            ]
        }
        workflow_config = {
            "max_queries": config["app"]["max_queries_dev"],
            "search_limit": config["app"]["search_limit_dev"],
            "auto_limit": config["app"]["search_auto_limit_dev"],
        }
    elif fast_mode:
        model_id = config["models"]["performance_low"]
        model_config = {
            key: model_id
            for key in [
                "create_queries",
                "check_relevance",
                "analyze_documents",
                "reflect_task",
                "final_report",
            ]
        }
        workflow_config = {
            "max_queries": config["app"]["max_queries_fast"],
            "search_limit": config["app"]["search_limit_fast"],
            "auto_limit": config["app"]["search_auto_limit_fast"],
        }
    else:
        model_config = {
            "create_queries": config["models"]["performance_low"],
            "check_relevance": config["models"]["performance_low"],
            "analyze_documents": config["models"]["performance_medium"],
            "reflect_task": config["models"]["performance_medium"],
            "final_report": config["models"]["performance_high"],
        }
        workflow_config = {
            "max_queries": config["app"]["max_queries"],
            "search_limit": config["app"]["search_limit"],
            "auto_limit": config["app"]["search_auto_limit"],
        }

    return model_config, workflow_config


def call_function_in_parallel(
    prompt_list: List[str],
    llm_function: Callable,
    max_workers: int = None,
    **llm_kwargs,
) -> List[Any]:
    """
    Run prompts in parallel using the given LLM function.

    Args:
        prompt_list (List[str]): List of prompts.
        llm_function (Callable): LLM function to call.
        max_workers (int, optional): Max worker threads.
        **llm_kwargs: Extra arguments for llm_function.

    Returns:
        List[Any]: Results for each prompt.
    """
    max_workers = max_workers or config["parallelization"]["max_workers"]
    results = [None] * len(prompt_list)

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        future_to_index = {
            executor.submit(llm_function, prompt, **llm_kwargs): i
            for i, prompt in enumerate(prompt_list)
        }

        for future in tqdm(
            concurrent.futures.as_completed(future_to_index),
            total=len(prompt_list),
            desc="Processing LLM queries in parallel...",
        ):
            index = future_to_index[future]
            try:
                results[index] = future.result()
            except Exception as exc:
                results[index] = f"Error: {exc}"

    return results


def create_docx_from_markdown(user_query, markdown_text):
    """
    Create a DOCX document from Markdown text and a user query.

    Args:
        user_query (str): Research question.
        markdown_text (str): Markdown-formatted text.

    Returns:
        bytes: Binary DOCX content.
    """
    doc = Document()
    doc.add_heading("Recherche-Bericht", 0)
    doc.add_paragraph(f"Recherchefrage: {user_query}")
    doc.add_paragraph(f"Erstellt am: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p = doc.add_paragraph()
    run = p.add_run(
        "Achtung: Dieser Bericht wurde mit einem KI-Recherche-Tool erstellt. Das Werkzeug ist experimentell. Ergebnisse können fehlerhaft oder unvollständig sein. Bitte prüfe die Ergebnisse immer."
    )
    run.bold = True
    run.font.color.rgb = RGBColor(255, 100, 100)

    # Split markdown into lines and process
    lines = markdown_text.split("\n")
    for line in lines:
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("##### "):
            doc.add_heading(line[6:].strip(), level=5)
        elif line.startswith("###### "):
            doc.add_heading(line[7:].strip(), level=6)
        elif re.match(r"^- ", line):
            # Level 1 list with dash
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, line[2:].strip())
        elif re.match(r"^\* ", line):
            # Level 1 list with asterisk
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, line[2:].strip())
        elif re.match(r"^  - ", line):
            # Level 2 list with dash (2 spaces)
            p = doc.add_paragraph(style="List Bullet 2")
            add_formatted_text(p, line[4:].strip())
        elif re.match(r"^    - ", line):
            # Level 2 list with dash (4 spaces)
            p = doc.add_paragraph(style="List Bullet 2")
            add_formatted_text(p, line[6:].strip())
        elif re.match(r"^  \* ", line):
            # Level 2 list with asterisk (2 spaces)
            p = doc.add_paragraph(style="List Bullet 2")
            add_formatted_text(p, line[4:].strip())
        elif re.match(r"^    \* ", line):
            # Level 2 list with asterisk (4 spaces)
            p = doc.add_paragraph(style="List Bullet 2")
            add_formatted_text(p, line[6:].strip())
        elif re.match(r"^      - ", line):
            # Level 3 list with dash (6 spaces)
            p = doc.add_paragraph(style="List Bullet 3")
            add_formatted_text(p, line[8:].strip())
        elif re.match(r"^        - ", line):
            # Level 3 list with dash (8 spaces)
            p = doc.add_paragraph(style="List Bullet 3")
            add_formatted_text(p, line[10:].strip())
        elif re.match(r"^      \* ", line):
            # Level 3 list with asterisk (6 spaces)
            p = doc.add_paragraph(style="List Bullet 3")
            add_formatted_text(p, line[8:].strip())
        elif re.match(r"^        \* ", line):
            # Level 3 list with asterisk (8 spaces)
            p = doc.add_paragraph(style="List Bullet 3")
            add_formatted_text(p, line[10:].strip())
        elif line and not line.startswith("#"):
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line.strip())

    # Save to BytesIO
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()


def add_formatted_text(paragraph, text):
    """
    Add formatted text to a Word paragraph, supporting bold, italic, and hyperlinks.

    Args:
        paragraph (docx.text.paragraph.Paragraph): Paragraph object.
        text (str): Text with Markdown-like formatting.
    """
    # Pattern to match bold (**text**), italic (*text*), and links ([text](url))
    pattern = r"(\*\*.*?\*\*|\*.*?\*|\[.*?\]\(.*?\))"

    parts = re.split(pattern, text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            # Markdown link [text](url)
            link_match = re.match(r"\[(.*?)\]\((.*?)\)", part)
            if link_match:
                link_text = link_match.group(1)
                link_url = link_match.group(2)

                # Add hyperlink to paragraph
                try:
                    # Create hyperlink
                    hyperlink = paragraph.part.document.part.relate_to(
                        link_url,
                        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                        is_external=True,
                    )

                    # Create hyperlink XML
                    hyperlink_xml = f'<w:hyperlink r:id="{hyperlink}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"><w:r><w:rPr><w:color w:val="0000FF"/><w:u w:val="single"/></w:rPr><w:t>{link_text}</w:t></w:r></w:hyperlink>'

                    # Add hyperlink to paragraph
                    hyperlink_element = parse_xml(hyperlink_xml)
                    paragraph._element.append(hyperlink_element)
                except Exception:
                    # Fallback: just add the link text with URL in parentheses
                    run = paragraph.add_run(f"{link_text} ({link_url})")
                    run.font.color.rgb = RGBColor(0, 0, 255)
        else:
            # Regular text
            if part:
                paragraph.add_run(part)
